import PyPDF2
import docx
import re
from typing import Optional, List, Dict, Any
import pdfplumber
from io import BytesIO
from intelligent_education_detector import IntelligentEducationDetector

class ImprovedResumeParser:
    def __init__(self):
        self.education_detector = IntelligentEducationDetector()
        self.skill_patterns = {
            'programming_languages': [
                'python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'php', 'go', 'rust',
                'typescript', 'kotlin', 'swift', 'scala', 'r', 'matlab', 'perl', 'c',
                'objective-c', 'dart', 'elixir', 'haskell', 'lua', 'shell', 'bash',
                'powershell', 'vb.net', 'visual basic', 'cobol', 'fortran', 'assembly'
            ],
            'web_technologies': [
                'html', 'css', 'react', 'angular', 'vue.js', 'vue', 'node.js', 'nodejs',
                'express', 'django', 'flask', 'spring', 'laravel', 'rails', 'asp.net',
                'jquery', 'bootstrap', 'sass', 'less', 'webpack', 'gulp', 'grunt',
                'next.js', 'nuxt.js', 'svelte', 'ember.js', 'backbone.js', 'redux',
                'mobx', 'styled-components', 'material-ui', 'tailwind', 'bulma'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
                'oracle', 'sqlite', 'dynamodb', 'firebase', 'mariadb', 'couchdb',
                'neo4j', 'influxdb', 'amazon rds', 'sql server', 'db2', 'snowflake', 'sql'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean',
                'docker', 'kubernetes', 'terraform', 'ansible', 'jenkins', 'gitlab ci',
                'github actions', 'circleci', 'travis ci', 'vagrant', 'puppet', 'chef'
            ],
            'data_science': [
                'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
                'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'tableau', 'power bi',
                'jupyter', 'anaconda', 'r studio', 'spss', 'sas', 'hadoop', 'spark',
                'kafka', 'airflow', 'mlflow', 'kubeflow', 'databricks', 'excel'
            ],
            'mobile_development': [
                'android', 'ios', 'react native', 'flutter', 'xamarin', 'cordova',
                'phonegap', 'ionic', 'unity', 'unreal engine'
            ],
            'tools_and_platforms': [
                'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
                'trello', 'asana', 'notion', 'figma', 'sketch', 'adobe xd', 'photoshop',
                'illustrator', 'vs code', 'intellij', 'eclipse', 'vim', 'emacs'
            ],
            'soft_skills': [
                'customer service', 'communication', 'leadership', 'teamwork', 'problem solving',
                'time management', 'organization', 'project management', 'public speaking',
                'negotiation', 'sales', 'marketing', 'training', 'mentoring'
            ],
            'business_skills': [
                'accounting', 'finance', 'economics', 'marketing', 'sales', 'hr',
                'operations', 'supply chain', 'logistics', 'consulting', 'strategy',
                'business analysis', 'data analysis', 'reporting', 'budgeting'
            ]
        }
        self.name_patterns = [
            r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\s*$',
            r'^([A-Z][a-z]+\s[A-Z]\.\s[A-Z][a-z]+)\s*$',
            r'^([A-Z][A-Z\s]+[A-Z])\s*$',
            r'Name:\s*([A-Z][a-z\s]+)',
            r'RESUME OF\s+([A-Z][a-z\s]+)',
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)(?:\s*[\n|])',
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)',
        ]
        self.phone_patterns = [
            r'(\+\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,6})',
            r'(\+\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,6})',
            r'(\d{10,15})',
            r'(\d{3,5}[-.\s]\d{3,4}[-.\s]\d{3,6})',
            r'(\d{2,4}[-.\s]\d{3,4}[-.\s]\d{3,4}[-.\s]\d{2,4})',
            r'(\(\d{2,4}\)[-.\s]?\d{3,4}[-.\s]?\d{3,6})',
            r'(\d{2,4}[-.\s]?\(\d{2,4}\)[-.\s]?\d{3,6})',
            r'(?:phone|mobile|tel|cell|contact)[:.]?\s*(\+?[\d\s\-\(\)\.]{8,20})',
            r'(\d{4,5}\s\d{3}\s\d{3,4})',
            r'(\d{2,3}\s\d{4}\s\d{4})',
            r'(\d{3}\s\d{3}\s\d{4})',
            r'(\d{2}[-.\s]\d{2}[-.\s]\d{2}[-.\s]\d{2}[-.\s]\d{2})',
            r'(\d{3}[-.\s]\d{3}[-.\s]\d{2}[-.\s]\d{2})',
            r'(\d{3}[-.\s]\d{4}[-.\s]\d{4})',
            r'(\d{2}[-.\s]\d{4}[-.\s]\d{4})',
            r'(\+91[-.\s]\d{5}[-.\s]\d{5})',
            r'(\+44[-.\s]\d{4}[-.\s]\d{6})',
            r'(\+33[-.\s]\d[-.\s]\d{2}[-.\s]\d{2}[-.\s]\d{2}[-.\s]\d{2})',
            r'(\+49[-.\s]\d{3}[-.\s]\d{3}[-.\s]\d{4})',
            r'(\+86[-.\s]\d{3}[-.\s]\d{4}[-.\s]\d{4})',
            r'(\+81[-.\s]\d{2}[-.\s]\d{4}[-.\s]\d{4})',
            r'(\+7[-.\s]\d{3}[-.\s]\d{3}[-.\s]\d{2}[-.\s]\d{2})',
        ]
        self.section_headers = {
            'summary': [
                'professional summary', 'summary', 'profile', 'objective', 'career objective',
                'professional profile', 'about me', 'overview', 'executive summary',
                'career summary', 'personal statement', 'professional overview'
            ],
            'experience': [
                'work experience', 'professional experience', 'employment history',
                'experience', 'career history', 'work history', 'employment',
                'work background', 'professional background'
            ],
            'skills': [
                'technical skills', 'skills', 'core competencies', 'technical competencies',
                'programming skills', 'technologies', 'expertise', 'proficiencies',
                'technical proficiencies', 'core skills', 'key skills', 'abilities'
            ],
            'education': [
                'education', 'academic background', 'qualifications', 'academic qualifications',
                'educational background', 'degrees', 'certifications', 'academic history',
                'university', 'college', 'school'
            ]
        }

    def extract_text(self, file_path: str, content_type: str, max_pages: int = 10) -> str:
        try:
            if content_type == 'application/pdf':
                return self._extract_from_pdf(file_path, max_pages)
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError("Unsupported file type")
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")

    def _extract_from_pdf(self, file_path: str, max_pages: int = 10) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_to_process = min(len(pdf.pages), max_pages)
                for i in range(pages_to_process):
                    page = pdf.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " ".join([cell for cell in row if cell]) + "\n"
            if text.strip():
                return self._clean_text(text)
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages_to_process = min(len(pdf_reader.pages), max_pages)
                for i in range(pages_to_process):
                    page = pdf_reader.pages[i]
                    text += page.extract_text() + "\n"
            if text.strip():
                return self._clean_text(text)
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
        raise Exception("Failed to extract text from PDF using all available methods")

    def _extract_from_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'[^\w\s@.,()|\-/+:]', ' ', text)
        return text.strip()

    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        name = self._extract_name_improved(text)
        email = self._extract_email_improved(text)
        phone = self._extract_phone_improved(text)
        location = self._extract_location_improved(text)
        return {
            'name': name,
            'email': email,
            'phone': phone,
            'location': location
        }

    def _extract_name_improved(self, text: str) -> Optional[str]:
        lines = text.split('\n')
        for line in lines[:10]:
            line = line.strip()
            if not line or len(line) < 3:
                continue
            skip_indicators = ['page', 'tip:', 'address', '@', 'phone', 'mobile', 'email']
            if any(indicator in line.lower() for indicator in skip_indicators):
                continue
            words = line.split()
            if len(words) >= 2 and len(words) <= 4:
                if all(word[0].isupper() and word[1:].islower() for word in words if len(word) > 1):
                    potential_name = ' '.join(words)
                    if self._validate_name_improved(potential_name):
                        return potential_name
        for pattern in self.name_patterns:
            for line in lines[:8]:
                match = re.search(pattern, line)
                if match:
                    potential_name = match.group(1).strip()
                    if self._validate_name_improved(potential_name):
                        return potential_name
        return None

    def _validate_name_improved(self, name: str) -> bool:
        if not name or len(name) < 3:
            return False
        exclusions = [
            'resume', 'curriculum vitae', 'cv', 'email', 'phone', 'address',
            'page', 'tip', 'career', 'objective', 'education', 'experience',
            'skills', 'work', 'contact', 'details', 'information'
        ]
        if any(exclusion in name.lower() for exclusion in exclusions):
            return False
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            return False
        letter_ratio = sum(c.isalpha() for c in name) / len(name)
        if letter_ratio < 0.7:
            return False
        return True

    def _extract_phone_improved(self, text: str) -> Optional[str]:
        for pattern in self.phone_patterns:
            matches = re.findall(pattern, text)
            for phone in matches:
                cleaned = self._clean_phone_improved(phone)
                if cleaned:
                    return cleaned
        return None

    def _clean_phone_improved(self, phone: str) -> Optional[str]:
        if not phone:
            return None
        digits_only = re.sub(r'[^\d+]', '', phone)
        digit_count = len(digits_only.replace('+', ''))
        if digit_count < 7 or digit_count > 15:
            return None
        cleaned_phone = phone.strip()
        cleaned_phone = re.sub(r'\s+', ' ', cleaned_phone)
        cleaned_phone = re.sub(r'[-.\s]{2,}', '-', cleaned_phone)
        if cleaned_phone.startswith('+'):
            return self._format_international_phone(cleaned_phone)
        elif digit_count >= 10:
            return self._format_domestic_phone(cleaned_phone, digit_count)
        else:
            return cleaned_phone

    def _format_international_phone(self, phone: str) -> str:
        digits = re.sub(r'[^\d]', '', phone)
        if len(digits) == 11 and digits.startswith('1'):
            return f"+1 {digits[1:4]} {digits[4:7]} {digits[7:]}"
        elif len(digits) == 12 and digits.startswith('61'):
            return f"+61 {digits[2]} {digits[3:7]} {digits[7:]}"
        elif len(digits) == 12 and digits.startswith('44'):
            return f"+44 {digits[2:6]} {digits[6:]}"
        elif len(digits) == 12 and digits.startswith('91'):
            return f"+91 {digits[2:7]} {digits[7:]}"
        elif len(digits) == 13 and digits.startswith('33'):
            return f"+33 {digits[2]} {digits[3:5]} {digits[5:7]} {digits[7:9]} {digits[9:]}"
        elif len(digits) == 12 and digits.startswith('49'):
            return f"+49 {digits[2:5]} {digits[5:8]} {digits[8:]}"
        elif len(digits) == 13 and digits.startswith('86'):
            return f"+86 {digits[2:5]} {digits[5:9]} {digits[9:]}"
        elif len(digits) == 12 and digits.startswith('81'):
            return f"+81 {digits[2:4]} {digits[4:8]} {digits[8:]}"
        else:
            return phone

    def _format_domestic_phone(self, phone: str, digit_count: int) -> str:
        digits = re.sub(r'[^\d]', '', phone)
        if digit_count == 10:
            return f"{digits[:3]} {digits[3:6]} {digits[6:]}"
        elif digit_count == 11:
            return f"{digits[0]} {digits[1:4]} {digits[4:7]} {digits[7:]}"
        elif digit_count == 9:
            return f"{digits[:3]} {digits[3:6]} {digits[6:]}"
        elif digit_count == 8:
            return f"{digits[:4]} {digits[4:]}"
        else:
            return phone.strip()

    def _extract_email_improved(self, text: str) -> Optional[str]:
        email_patterns = [
            r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b',
            r'Email:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            r'E-mail:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            r'Contact:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        ]
        for pattern in email_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for email in matches:
                if self._validate_email_improved(email):
                    return email.lower()
        return None

    def _validate_email_improved(self, email: str) -> bool:
        if not email or '@' not in email:
            return False
        parts = email.split('@')
        if len(parts) != 2:
            return False
        local, domain = parts
        if not local or not domain or '.' not in domain:
            return False
        if len(local) < 1 or len(domain) < 3:
            return False
        return True

    def _extract_location_improved(self, text: str) -> Optional[str]:
        location_patterns = [
            r'([A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+,\s*\d{4})',
            r'([A-Z][a-z]+,\s*[A-Z]{2,3}(?:\s+\d{4,5})?)',
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',
            r'(\d+\s+[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+)',
            r'Location:\s*([A-Z][a-zA-Z\s,.-]+)',
            r'Address:\s*([A-Z][a-zA-Z\s,.-]+)',
            r'Based in:\s*([A-Z][a-zA-Z\s,.-]+)',
        ]
        lines = text.split('\n')
        for line in lines[:5]:
            for pattern in location_patterns:
                matches = re.findall(pattern, line)
                for location in matches:
                    if self._validate_location_improved(location):
                        return location.strip()
        return None

    def _validate_location_improved(self, location: str) -> bool:
        if not location or len(location) < 3:
            return False
        if len(location) > 100:
            return False
        location_lower = location.lower()
        exclude_words = ['email', 'phone', 'mobile', 'resume', 'page']
        if any(word in location_lower for word in exclude_words):
            return False
        return True

    def extract_skills(self, text: str) -> List[str]:
        skills = set()
        text_lower = text.lower()
        for category, skill_list in self.skill_patterns.items():
            for skill in skill_list:
                if f' {skill.lower()} ' in f' {text_lower} ':
                    skills.add(skill.title())
                elif re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
                    skills.add(skill.title())
        skills_section = self._extract_section(text, 'skills')
        if skills_section:
            additional_skills = self._parse_skills_from_section(skills_section)
            skills.update(additional_skills)
        return list(skills)[:25]

    def _parse_skills_from_section(self, section: str) -> set:
        skills = set()
        lines = section.split('\n')
        for line in lines:
            if 'tip:' in line.lower():
                continue
            line = re.sub(r'^[•\-\*]\s*', '', line.strip())
            for category, skill_list in self.skill_patterns.items():
                for skill in skill_list:
                    if re.search(r'\b' + re.escape(skill.lower()) + r'\b', line.lower()):
                        skills.add(skill.title())
        return skills

    def extract_education_info(self, text: str) -> Dict[str, Any]:
        education_details = self.education_detector.detect_with_details(text)
        education_info = {
            'institutions': [inst['name'] for inst in education_details['institutions']],
            'best_institution': education_details['best_match'],
            'degrees': [],
            'fields': [],
            'graduation_years': [],
            'detection_confidence': education_details['institutions'][0]['confidence'] if education_details['institutions'] else 0,
            'context_analysis': education_details['analysis']
        }
        degree_patterns = [
            r'(B\.?Tech\.?|Bachelor.*?Technology|BTech)',
            r'(B\.?E\.?|Bachelor.*?Engineering|BE)',
            r'(B\.?Sc\.?|Bachelor.*?Science|BSc)',
            r'(B\.?A\.?|Bachelor.*?Arts|BA)',
            r'(B\.?Com\.?|Bachelor.*?Commerce|BCom)',
            r'(M\.?Tech\.?|Master.*?Technology|MTech)',
            r'(M\.?Sc\.?|Master.*?Science|MSc)',
            r'(M\.?A\.?|Master.*?Arts|MA)',
            r'(MBA|Master.*?Business)',
            r'(PhD|Ph\.?D\.?|Doctorate)',
            r'(Diploma|Certificate)',
        ]
        field_patterns = [
            r'B\.?Tech\.?\s+in\s+([^.\n,]+)',
            r'Bachelor.*?in\s+([^.\n,]+)',
            r'Master.*?in\s+([^.\n,]+)',
            r'PhD.*?in\s+([^.\n,]+)',
            r'(Computer Science.*?Engineering|CSE)',
            r'(Computer Science)',
            r'(Information Technology|IT)',
            r'(Electronics.*?Communication|ECE)',
            r'(Mechanical Engineering)',
            r'(Civil Engineering)',
            r'(Electrical Engineering)',
            r'(Business Administration)',
            r'(Data Science)',
            r'(Artificial Intelligence)',
        ]
        year_patterns = [
            r'(20\d{2})\s*[-–]\s*(20\d{2})',
            r'(20\d{2})\s*[-–]\s*present',
            r'graduating.*?(20\d{2})',
            r'batch.*?(20\d{2})',
        ]
        lines = text.split('\n')
        for line in lines:
            for pattern in degree_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    if match not in education_info['degrees']:
                        education_info['degrees'].append(match)
            for pattern in field_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    if len(match) > 3 and match not in education_info['fields']:
                        field_clean = re.sub(r'\s+', ' ', match).strip()
                        education_info['fields'].append(field_clean)
            for pattern in year_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        for year in match:
                            if year.isdigit() and 1950 <= int(year) <= 2030:
                                if year not in education_info['graduation_years']:
                                    education_info['graduation_years'].append(year)
                    elif match.isdigit() and 1950 <= int(match) <= 2030:
                        if match not in education_info['graduation_years']:
                            education_info['graduation_years'].append(match)
        return education_info

    def extract_professional_summary(self, text: str) -> Optional[str]:
        summary_section = self._extract_section(text, 'summary')
        if summary_section:
            summary = re.sub(r'\n+', ' ', summary_section).strip()
            summary = re.sub(r'\(Tip:.*?\)', '', summary, flags=re.DOTALL)
            summary = summary.strip()
            if 50 < len(summary) < 1000:
                return summary
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs[1:4]:
            clean_para = paragraph.strip().replace('\n', ' ')
            clean_para = re.sub(r'\(Tip:.*?\)', '', clean_para, flags=re.DOTALL)
            clean_para = clean_para.strip()
            objective_keywords = ['seeking', 'looking', 'aim', 'objective', 'goal', 'enthusiast']
            if (len(clean_para) > 50 and 
                any(keyword in clean_para.lower() for keyword in objective_keywords)):
                return clean_para[:500]
        return "Professional summary not found in resume."

    def _extract_section(self, text: str, section_type: str) -> Optional[str]:
        if section_type not in self.section_headers:
            return None
        headers = self.section_headers[section_type]
        text_lines = text.split('\n')
        for i, line in enumerate(text_lines):
            line_lower = line.lower().strip()
            for header in headers:
                if (line_lower == header or 
                    line_lower.startswith(header + ':') or
                    header in line_lower and len(line_lower) < len(header) + 15):
                    section_content = []
                    for j in range(i + 1, len(text_lines)):
                        next_line = text_lines[j].strip()
                        if self._is_section_header(next_line.lower()):
                            break
                        if next_line:
                            section_content.append(next_line)
                        if len(section_content) > 25:
                            break
                    if section_content:
                        return '\n'.join(section_content)
        return None

    def _is_section_header(self, line: str) -> bool:
        section_indicators = [
            'education', 'experience', 'work', 'skills', 'objective', 'summary',
            'employment', 'career', 'background', 'qualifications', 'achievements',
            'projects', 'certifications', 'awards', 'interests', 'hobbies',
            'references', 'leadership', 'volunteer', 'activities'
        ]
        line_clean = line.strip().lower()
        for indicator in section_indicators:
            if (line_clean == indicator or 
                line_clean.startswith(indicator + ':') or
                line_clean.startswith(indicator + ' ') or
                line_clean.endswith(indicator)):
                return True
        return False
